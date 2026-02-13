import os
from flask import render_template
from xhtml2pdf import pisa
from docx import Document
from io import BytesIO
import pandas as pd
from datetime import datetime

class DocumentEngine:
    @staticmethod
    def generate_pdf(template_name, context, output_path=None):
        """Generates a PDF from a Jinja2 template."""
        html = render_template(template_name, **context)
        result = BytesIO()
        pisa_status = pisa.CreatePDF(html, dest=result)
        
        if pisa_status.err:
            return None
        
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(result.getvalue())
            return output_path
        
        return result.getvalue()

    @staticmethod
    def generate_word(template_data, output_path):
        """
        Generates a Word document using placeholders.
        template_data: dict of {placeholder: value}
        """
        # Note: This is a simplified version. For a production system, 
        # you'd use a .docx template and find/replace.
        doc = Document()
        doc.add_heading(template_data.get('title', 'Document'), 0)
        
        for key, value in template_data.items():
            if key != 'title':
                doc.add_paragraph(f"{key}: {value}")
        
        doc.save(output_path)
        return output_path

    @staticmethod
    def generate_excel(data_list, columns, output_path):
        """Generates an Excel file from a list of dictionaries."""
        df = pd.DataFrame(data_list, columns=columns)
        df.to_excel(output_path, index=False)
        return output_path

def log_document(doc_type, doc_name, file_path, status='FINAL', fy=None):
    from models import db, DocumentLog
    new_doc = DocumentLog(
        doc_type=doc_type,
        doc_name=doc_name,
        file_path=file_path,
        status=status,
        financial_year=fy or f"{datetime.now().year}"
    )
    db.session.add(new_doc)
    db.session.commit()
    return new_doc
